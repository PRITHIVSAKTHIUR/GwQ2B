import gradio as gr
import requests
import uuid 
from PIL import Image
import os
import io
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Image as RLImage, Paragraph, Spacer
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Define model options
MODEL_OPTIONS = {
    "Qwen2VL Base": "Qwen/Qwen2-VL-2B-Instruct",
    "Latex OCR": "prithivMLmods/Qwen2-VL-OCR-2B-Instruct",
    "Math Prase": "prithivMLmods/Qwen2-VL-Math-Prase-2B-Instruct",
    "Text Analogy Ocrtest": "prithivMLmods/Qwen2-VL-Ocrtest-2B-Instruct"
}

# Gradio API endpoint
GRADIO_API_URL = "https://prithivmlmods-qwen2-vl-2b.hf.space/"

def identify_and_save_blob(blob_path):
    """Identifies if the blob is an image and saves it."""
    try:
        with open(blob_path, 'rb') as file:
            blob_content = file.read()
            try:
                Image.open(io.BytesIO(blob_content)).verify()  # Check if it's a valid image
                extension = ".png"  # Default to PNG for saving
                media_type = "image"
            except (IOError, SyntaxError):
                raise ValueError("Unsupported media type. Please upload a valid image.")

            filename = f"temp_{uuid.uuid4()}_media{extension}"
            with open(filename, "wb") as f:
                f.write(blob_content)

            return filename, media_type

    except FileNotFoundError:
        raise ValueError(f"The file {blob_path} was not found.")
    except Exception as e:
        raise ValueError(f"An error occurred while processing the file: {e}")

def qwen_inference(model_name, media_input, text_input=None):
    """Handles inference using the Gradio API."""
    if isinstance(media_input, str):
        media_path = media_input
        if not media_path.endswith(('.png', '.jpg', '.jpeg', '.bmp', '.gif')):
            try:
                media_path, media_type = identify_and_save_blob(media_input)
            except Exception as e:
                raise ValueError("Unsupported media type. Please upload a valid image.")
    else:
        raise ValueError("Invalid media input.")

    # Prepare the API request
    files = {'media_input': open(media_path, 'rb')}
    data = {
        'model_name': model_name,
        'text_input': text_input,
        'api_name': '/qwen_inference'
    }

    # Send the request to the Gradio API
    response = requests.post(GRADIO_API_URL, files=files, data=data)
    if response.status_code == 200:
        return response.json()['output']
    else:
        raise ValueError(f"API request failed with status code {response.status_code}: {response.text}")

def format_plain_text(output_text):
    """Formats the output text as plain text without LaTeX delimiters."""
    plain_text = output_text.replace("\\(", "").replace("\\)", "").replace("\\[", "").replace("\\]", "")
    return plain_text

def generate_document(media_path, output_text, file_format, font_choice, font_size, line_spacing, alignment, image_size):
    """Generates a document with the input image and plain text output."""
    plain_text = format_plain_text(output_text)
    if file_format == "pdf":
        return generate_pdf(media_path, plain_text, font_choice, font_size, line_spacing, alignment, image_size)
    elif file_format == "docx":
        return generate_docx(media_path, plain_text, font_choice, font_size, line_spacing, alignment, image_size)

def generate_pdf(media_path, plain_text, font_choice, font_size, line_spacing, alignment, image_size):
    """Generates a PDF document."""
    filename = f"output_{uuid.uuid4()}.pdf"
    doc = SimpleDocTemplate(
        filename,
        pagesize=A4,
        rightMargin=inch,
        leftMargin=inch,
        topMargin=inch,
        bottomMargin=inch
    )
    styles = getSampleStyleSheet()
    styles["Normal"].fontName = font_choice
    styles["Normal"].fontSize = int(font_size)
    styles["Normal"].leading = int(font_size) * line_spacing
    styles["Normal"].alignment = {
        "Left": 0,
        "Center": 1,
        "Right": 2,
        "Justified": 4
    }[alignment]

    # Register font
    font_path = f"font/{font_choice}"
    pdfmetrics.registerFont(TTFont(font_choice, font_path))

    story = []

    # Add image with size adjustment
    image_sizes = {
        "Small": (200, 200),
        "Medium": (400, 400),
        "Large": (600, 600)
    }
    img = RLImage(media_path, width=image_sizes[image_size][0], height=image_sizes[image_size][1])
    story.append(img)
    story.append(Spacer(1, 12))

    # Add plain text output
    text = Paragraph(plain_text, styles["Normal"])
    story.append(text)

    doc.build(story)
    return filename

def generate_docx(media_path, plain_text, font_choice, font_size, line_spacing, alignment, image_size):
    """Generates a DOCX document."""
    filename = f"output_{uuid.uuid4()}.docx"
    doc = docx.Document()

    # Add image with size adjustment
    image_sizes = {
        "Small": docx.shared.Inches(2),
        "Medium": docx.shared.Inches(4),
        "Large": docx.shared.Inches(6)
    }
    doc.add_picture(media_path, width=image_sizes[image_size])
    doc.add_paragraph()

    # Add plain text output
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.line_spacing = line_spacing
    paragraph.paragraph_format.alignment = {
        "Left": WD_ALIGN_PARAGRAPH.LEFT,
        "Center": WD_ALIGN_PARAGRAPH.CENTER,
        "Right": WD_ALIGN_PARAGRAPH.RIGHT,
        "Justified": WD_ALIGN_PARAGRAPH.JUSTIFY
    }[alignment]
    run = paragraph.add_run(plain_text)
    run.font.name = font_choice
    run.font.size = docx.shared.Pt(int(font_size))

    doc.save(filename)
    return filename

# CSS for output styling
css = """
  #output {
    height: 500px; 
    overflow: auto; 
    border: 1px solid #ccc; 
  }
.submit-btn {
    background-color: #cf3434  !important;
    color: white !important;
}
.submit-btn:hover {
    background-color: #ff2323 !important;
}
.download-btn {
    background-color: #35a6d6 !important;
    color: white !important;
}
.download-btn:hover {
    background-color: #22bcff !important;
}
"""

# Gradio app setup
with gr.Blocks(css=css) as demo:
    gr.Markdown("# Qwen2VL Models: Vision and Language Processing")

    with gr.Tab(label="Image Input"):
        with gr.Row():
            with gr.Column():
                model_choice = gr.Dropdown(
                    label="Model Selection",
                    choices=list(MODEL_OPTIONS.keys()),
                    value="Latex OCR"
                )
                input_media = gr.File(
                    label="Upload Image", type="filepath" 
                )
                text_input = gr.Textbox(label="Question", placeholder="Ask a question about the image...")
                submit_btn = gr.Button(value="Submit", elem_classes="submit-btn")

            with gr.Column():
                output_text = gr.Textbox(label="Output Text", lines=10)
                plain_text_output = gr.Textbox(label="Standardized Plain Text", lines=10)

        submit_btn.click(
            qwen_inference, [model_choice, input_media, text_input], [output_text]
        ).then(
            lambda output_text: format_plain_text(output_text), [output_text], [plain_text_output]
        )

        # Add examples directly usable by clicking
        with gr.Row():
            gr.Examples(
                examples=[
                    ["examples/1.png", "summarize the letter", "Text Analogy Ocrtest"],
                    ["examples/2.jpg", "Summarize the full image in detail", "Latex OCR"],
                    ["examples/3.png", "Describe the photo", "Qwen2VL Base"],
                    ["examples/4.png", "summarize and solve the problem", "Math Prase"],
                ],
                inputs=[input_media, text_input, model_choice],
                outputs=[output_text, plain_text_output],
                fn=lambda img, question, model: qwen_inference(model, img, question),
                cache_examples=False,
            )

        with gr.Row():
            with gr.Column():
                line_spacing = gr.Dropdown(
                    choices=[0.5, 1.0, 1.15, 1.5, 2.0, 2.5, 3.0],
                    value=1.5,
                    label="Line Spacing"
                )
                font_size = gr.Dropdown(
                    choices=["8", "10", "12", "14", "16", "18", "20", "22", "24"],
                    value="18",
                    label="Font Size"
                )
                font_choice = gr.Dropdown(
                    choices=[
                        "DejaVuMathTeXGyre.ttf", 
                        "FiraCode-Medium.ttf", 
                        "InputMono-Light.ttf",
                        "JetBrainsMono-Thin.ttf", 
                        "ProggyCrossed Regular Mac.ttf", 
                        "SourceCodePro-Black.ttf", 
                        "arial.ttf", 
                        "calibri.ttf", 
                        "mukta-malar-extralight.ttf", 
                        "noto-sans-arabic-medium.ttf", 
                        "times new roman.ttf",
                        "ANGSA.ttf",
                        "Book-Antiqua.ttf",
                        "CONSOLA.TTF",
                        "COOPBL.TTF",
                        "Rockwell-Bold.ttf",
                        "Candara Light.TTF",
                        "Carlito-Regular.ttf Carlito-Regular.ttf",
                        "Castellar.ttf",
                        "Courier New.ttf",
                        "LSANS.TTF",
                        "Lucida Bright Regular.ttf",
                        "TRTempusSansITC.ttf",
                        "Verdana.ttf",
                        "bell-mt.ttf",
                        "eras-itc-light.ttf",
                        "fonnts.com-aptos-light.ttf",
                        "georgia.ttf",
                        "segoeuithis.ttf",
                        "youyuan.TTF",
                        "TfPonetoneExpanded-7BJZA.ttf",
                    ],
                    value="youyuan.TTF",
                    label="Font Choice"
                )
                alignment = gr.Dropdown(
                    choices=["Left", "Center", "Right", "Justified"],
                    value="Justified",
                    label="Text Alignment"
                )
                image_size = gr.Dropdown(
                    choices=["Small", "Medium", "Large"],
                    value="Small",
                    label="Image Size"
                )
                file_format = gr.Radio(["pdf", "docx"], label="File Format", value="pdf")
                get_document_btn = gr.Button(value="Get Document", elem_classes="download-btn")

        get_document_btn.click(
            generate_document, [input_media, output_text, file_format, font_choice, font_size, line_spacing, alignment, image_size], gr.File(label="Download Document")
        )

demo.launch(debug=True)